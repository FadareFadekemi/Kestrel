"""
CV export: parse plain-text CV → DOCX (python-docx) or PDF (fpdf2).
Both outputs use Times New Roman 12pt, single-column, ATS-compliant.
"""

import io
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fpdf import FPDF


# ── Section header registry ───────────────────────────────────────────────────

_HEADERS = {
    "PROFESSIONAL SUMMARY", "SUMMARY", "PROFESSIONAL PROFILE", "PROFILE",
    "SKILLS", "CORE COMPETENCIES", "TECHNICAL SKILLS", "KEY SKILLS", "CORE SKILLS",
    "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EXPERIENCE", "EMPLOYMENT HISTORY",
    "EDUCATION", "ACADEMIC BACKGROUND", "EDUCATIONAL BACKGROUND",
    "CERTIFICATIONS", "CERTIFICATES", "PROFESSIONAL CERTIFICATIONS", "TRAINING",
    "PROJECTS", "KEY PROJECTS",
    "LANGUAGES", "ADDITIONAL INFORMATION", "VOLUNTEER EXPERIENCE", "REFERENCES",
}

_SECTION_ORDER = [
    "PROFESSIONAL SUMMARY", "SUMMARY", "PROFILE",
    "SKILLS", "CORE COMPETENCIES", "TECHNICAL SKILLS", "KEY SKILLS",
    "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EXPERIENCE",
    "EDUCATION",
    "CERTIFICATIONS", "CERTIFICATES",
    "PROJECTS",
    "LANGUAGES",
    "ADDITIONAL INFORMATION",
]

_EXPERIENCE_KEYS = {"PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EXPERIENCE", "EMPLOYMENT HISTORY"}
_EDUCATION_KEYS  = {"EDUCATION", "ACADEMIC BACKGROUND", "EDUCATIONAL BACKGROUND"}


# ── Parser ────────────────────────────────────────────────────────────────────

def parse_cv_text(text: str) -> dict:
    """
    Returns:
        {"name": str, "contact": str, "sections": {HEADER: [lines]}}
    """
    result = {"name": "", "contact": "", "sections": {}}
    current = None
    buf: list[str] = []

    for raw in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line    = raw.rstrip()
        stripped = line.strip()
        upper   = stripped.upper()

        if upper in _HEADERS:
            if current is not None:
                result["sections"][current] = buf
            current = upper
            buf = []
        elif current is None:
            if not stripped:
                continue
            if not result["name"]:
                result["name"] = stripped
            elif not result["contact"]:
                result["contact"] = stripped
        else:
            buf.append(line)

    if current is not None:
        result["sections"][current] = buf

    return result


def _is_bullet(line: str) -> bool:
    return bool(re.match(r"^\s*[•\-\*]", line))


def _bullet_text(line: str) -> str:
    return re.sub(r"^\s*[•\-\*]\s*", "", line).strip()


def _is_pipe_entry(line: str) -> bool:
    return "|" in line and not _is_bullet(line)


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _font(run, size: float, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic


def _spacing(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    para.paragraph_format.line_spacing = Pt(14)


def _add_bottom_border(para, color="000000", sz="4"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _docx_section_header(doc, text: str):
    p = doc.add_paragraph()
    _spacing(p, before=10, after=2)
    run = p.add_run(text)
    _font(run, 12, bold=True)
    _add_bottom_border(p)
    return p


def _docx_pipe_entry(doc, line: str, bold_first=True, before=6):
    parts = [x.strip() for x in line.split("|")]
    p = doc.add_paragraph()
    _spacing(p, before=before, after=1)
    if parts:
        run = p.add_run(parts[0])
        _font(run, 11, bold=bold_first)
    if len(parts) > 1:
        run2 = p.add_run("  |  " + "  |  ".join(parts[1:]))
        _font(run2, 11)
    return p


def _docx_bullet(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    _spacing(p, before=1, after=1)
    run = p.add_run(f"•  {text}")
    _font(run, 11)
    return p


def _docx_body(doc, text: str):
    p = doc.add_paragraph()
    _spacing(p, before=2, after=2)
    run = p.add_run(text)
    _font(run, 11)
    return p


def _render_docx_section(doc, key: str, lines: list[str]):
    _docx_section_header(doc, key)
    non_empty = [l for l in lines if l.strip()]
    if not non_empty:
        return

    if key in _EXPERIENCE_KEYS:
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if _is_bullet(s):
                _docx_bullet(doc, _bullet_text(s))
            elif _is_pipe_entry(s):
                _docx_pipe_entry(doc, s, bold_first=True, before=6)
            else:
                p = doc.add_paragraph()
                _spacing(p, before=4, after=1)
                run = p.add_run(s)
                _font(run, 11, bold=True)

    elif key in _EDUCATION_KEYS:
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if _is_pipe_entry(s):
                _docx_pipe_entry(doc, s, bold_first=True, before=4)
            else:
                _docx_body(doc, s)

    else:
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if _is_bullet(s):
                _docx_bullet(doc, _bullet_text(s))
            else:
                _docx_body(doc, s)


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_docx(cv_text: str) -> bytes:
    parsed = parse_cv_text(cv_text)
    doc    = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    # Remove the default empty paragraph Word adds
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(name_para, before=0, after=2)
    run = name_para.add_run((parsed["name"] or "").upper())
    _font(run, 16, bold=True)

    # Contact
    if parsed["contact"]:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(cp, before=0, after=4)
        run = cp.add_run(parsed["contact"])
        _font(run, 11)

    # Divider under contact
    _add_bottom_border(doc.add_paragraph())

    # Sections in canonical order, then any leftovers
    seen = set()
    for key in _SECTION_ORDER:
        if key in parsed["sections"]:
            _render_docx_section(doc, key, parsed["sections"][key])
            seen.add(key)
    for key, lines in parsed["sections"].items():
        if key not in seen:
            _render_docx_section(doc, key, lines)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF helpers ───────────────────────────────────────────────────────────────

class _CVPDF(FPDF):
    def header(self): pass
    def footer(self): pass


def _pdf_section_header(pdf: _CVPDF, text: str):
    pdf.ln(5)
    pdf.set_font("Times", "B", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 6, text, ln=True)
    # underline rule
    x = pdf.get_x()
    y = pdf.get_y()
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.3)
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(2)


def _pdf_pipe_entry(pdf: _CVPDF, line: str, bold_first=True, top_gap=3):
    parts = [x.strip() for x in line.split("|")]
    pdf.ln(top_gap)
    if parts:
        if bold_first:
            pdf.set_font("Times", "B", 11)
            pdf.write(5, parts[0])
        else:
            pdf.set_font("Times", "", 11)
            pdf.write(5, parts[0])
    if len(parts) > 1:
        pdf.set_font("Times", "", 11)
        pdf.write(5, "  |  " + "  |  ".join(parts[1:]))
    pdf.ln(5)


def _pdf_bullet(pdf: _CVPDF, text: str):
    pdf.set_font("Times", "", 11)
    pdf.set_x(pdf.l_margin + 5)
    pdf.multi_cell(0, 5, f"-  {text}", ln=True)


def _pdf_body(pdf: _CVPDF, text: str):
    pdf.set_font("Times", "", 11)
    pdf.multi_cell(0, 5, text, ln=True)


def _render_pdf_section(pdf: _CVPDF, key: str, lines: list[str]):
    _pdf_section_header(pdf, key)
    non_empty = [l for l in lines if l.strip()]
    if not non_empty:
        return

    if key in _EXPERIENCE_KEYS:
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if _is_bullet(s):
                _pdf_bullet(pdf, _bullet_text(s))
            elif _is_pipe_entry(s):
                _pdf_pipe_entry(pdf, s, bold_first=True, top_gap=4)
            else:
                pdf.ln(3)
                pdf.set_font("Times", "B", 11)
                pdf.multi_cell(0, 5, s, ln=True)

    elif key in _EDUCATION_KEYS:
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if _is_pipe_entry(s):
                _pdf_pipe_entry(pdf, s, bold_first=True, top_gap=3)
            else:
                _pdf_body(pdf, s)

    else:
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if _is_bullet(s):
                _pdf_bullet(pdf, _bullet_text(s))
            else:
                _pdf_body(pdf, s)


# ── PDF builder ───────────────────────────────────────────────────────────────

def build_pdf(cv_text: str) -> bytes:
    parsed = parse_cv_text(cv_text)
    pdf    = _CVPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(25, 20, 25)
    pdf.set_text_color(0, 0, 0)

    # Name
    pdf.set_font("Times", "B", 16)
    pdf.cell(0, 9, (parsed["name"] or "").upper(), ln=True, align="C")

    # Contact
    if parsed["contact"]:
        pdf.set_font("Times", "", 10)
        pdf.cell(0, 5, parsed["contact"], ln=True, align="C")

    # Divider
    pdf.ln(2)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.4)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(3)

    # Sections
    seen = set()
    for key in _SECTION_ORDER:
        if key in parsed["sections"]:
            _render_pdf_section(pdf, key, parsed["sections"][key])
            seen.add(key)
    for key, lines in parsed["sections"].items():
        if key not in seen:
            _render_pdf_section(pdf, key, lines)

    return bytes(pdf.output())
